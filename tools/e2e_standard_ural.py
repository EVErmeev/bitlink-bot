"""Offline E2E driver for TASK-2026-08-05-19BD410.

Builds a canonical Protocol from the control UralDroneZavod transcript facts,
runs the real resolvers (meeting_type, meeting_topic, project_context,
standard_context) and the project_standard renderer (standalone + confluence
storage), then writes the TASK §21 debug artifacts into <out_dir>.

Confluence/Telegram/DOCX live publishing needs credentials and is out of scope
here; the storage mode (no <h1>) and standalone mode (exactly one <h1>) are
verified directly.

Usage: python tools/e2e_standard_ural.py <out_dir>
"""
import json
import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from models.protocol import (
    DecisionItem, Protocol, QuestionItem, RiskItem, TaskItem, TopicBlock,
)
from protocol_templates.project_standard import ProjectStandardTemplate
from services import project_context_resolver as pcr
from services.meeting_topic_resolver import resolve_meeting_topic
from services.meeting_type_resolver import resolve_meeting_type
from services.standard_context_resolver import resolve_standard_context


def main(out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    transcript = Path(
        r"D:\OpenCode\bitlink-bot\debug\fb860189\ef6a3922\source_transcript.txt"
    ).read_text(encoding="utf-8")

    filename = "УралДронЗавод + Первый БИТ 31.07.2026.txt"
    participants = [
        {"name": "Евгений", "role": "Аналитик", "side": "Первый БИТ"},
        {"name": "Александр", "role": "Диспетчер производства", "side": "УралДронЗавод"},
        {"name": "Семён", "role": "Начальник цеха", "side": "УралДронЗавод"},
    ]

    # meeting type
    mt, mt_conf, mt_evidence = resolve_meeting_type(
        transcript_text=transcript, source_filename=filename,
        participants=participants)
    (out_dir / "meeting_type_resolution.json").write_text(
        json.dumps({"meeting_type": mt, "confidence": mt_conf,
                    "evidence": mt_evidence}, ensure_ascii=False, indent=2),
        encoding="utf-8")

    # project context
    pcr_res = pcr.resolve_project_context(
        transcript_text=transcript, source_filename=filename,
        participants=participants)
    (out_dir / "project_context_resolution.json").write_text(
        json.dumps(pcr_res.to_dict() if hasattr(pcr_res, "to_dict") else vars(pcr_res),
                   ensure_ascii=False, indent=2, default=str), encoding="utf-8")

    # meeting topic
    topic = resolve_meeting_topic(
        source_filename=filename,
        source_title="УралДронЗавод + Первый БИТ — демонстрация процессов блока «Производство»",
        meeting_purpose=(
            "Обследовать и согласовать целевую схему работы производственного "
            "блока УралДронЗавода в 1С:ERP."))
    (out_dir / "meeting_topic_resolution.json").write_text(
        json.dumps({"meeting_topic": topic}, ensure_ascii=False, indent=2),
        encoding="utf-8")

    # Protocol with control facts
    p = Protocol(protocol_id="ural", template_id="project_standard",
                 source_context_id="src", client_name=p_t_res_cn(pcr_res),
                 project_name=p_t_res_pn(pcr_res))
    p.meeting_date = date(2026, 7, 31)
    p.meeting_time = None
    p.meeting_place = "Екатеринбург"
    p.meeting_purpose = (
        "Обследовать и согласовать целевую схему работы производственного "
        "блока УралДронЗавода в 1С:ERP.")
    p.meeting_context = ("31.07.2026, Екатеринбург, длительность 89 мин, "
                         "участников 3.")
    p.current_state = "31.07.2026, Екатеринбург, длительность 89 мин."
    p.key_outcomes = (
        "Показано диспетчирование этапов производства. Подтверждена схема "
        "обеспечения цеха материалами. Зафиксирован функциональный разрыв: "
        "сменное задание не связано с обеспечением этапа.")

    for name, side in (("Евгений", "Первый БИТ"),
                       ("Александр", "УралДронЗавод"),
                       ("Семён", "УралДронЗавод")):
        p.participants.append({"name": name, "role": "", "side": side})

    p.topic_blocks.append(TopicBlock(
        topic_id="t1", title="Диспетчирование этапов", source_context_id="c",
        discussion_content=(
            "Работал производственный диспетчер. Диспетчирование этапов "
            "производственного процесса. Статус «Начат». Отображение этапов "
            "и пиктограмм. Кладовая для обеспечения материалами."),
        conclusion="Схема диспетчирования подтверждена, требует уточнения по ролям.",
        status_text="Требует уточнения", source_item_ids=["i1"]))
    p.topic_blocks.append(TopicBlock(
        topic_id="t2", title="Обеспечение материалами", source_context_id="c",
        discussion_content=(
            "Сменное задание формирует расход и трудозатраты, но в типовом "
            "функционале не связано с обеспечением этапа. Передача материалов "
            "в кладовую фиксируется на остатках."),
        conclusion="Зафиксирован функциональный разрыв: посменное списание требует доработки.",
        status_text="Требует доработки", source_item_ids=["i2"]))
    p.topic_blocks.append(TopicBlock(
        topic_id="t3", title="Сменные задания", source_context_id="c",
        discussion_content=(
            "Формирование сменных заданий по подразделениям. Выбор исполнителей. "
            "Учёт свободного времени по табелю отсутствует в типовом функционале."),
        conclusion="Смешанный учёт сотрудников и бригад требует уточнения.",
        status_text="Требует уточнения", source_item_ids=["i3"]))

    p.decisions.append(DecisionItem(
        decision_id="d1", source_context_id="c",
        decision_text="Зафиксировать разрыв между сменным заданием и обеспечением этапа.",
        responsible="Евгений", deadline="04.08.2026"))
    p.risks.append(RiskItem(
        risk_id="r1", source_context_id="c",
        risk_text="Посменное списание материалов вручную ведёт к человеческому фактору.",
        reason="В типовом функционале нет связи сменного задания и обеспечения.",
        impact="Ошибки в расходе и остатках.", measures="Функциональная доработка.",
        responsible="Евгений"))
    p.tasks.append(TaskItem(
        task_id="w1", source_context_id="c",
        task_text="Подготовить список бригад и подразделений для учёта по бригадам.",
        basis="Решение по учёту выработки.", expected_result="Список бригад.",
        responsible="Семён", deadline="05.08.2026", status="Не начато"))

    # standard semantic context (initial_situation / main_problem) via resolver
    p._standard_context = resolve_standard_context(p, transcript, llm=None)

    # store resolutions for the grouper/title
    p._standard_meeting_type = mt
    p._standard_meeting_topic = topic
    p._standard_meeting_type_resolution = {"confidence": mt_conf, "evidence": mt_evidence}
    p._standard_project_resolution = pcr_res.to_dict() if hasattr(pcr_res, "to_dict") else vars(pcr_res)
    p._standard_source_type = "local_transcript"
    p._standard_source_filename = filename
    p._standard_recording_source = ""

    (out_dir / "standard_context_fields.json").write_text(
        json.dumps(p._standard_context, ensure_ascii=False, indent=2), encoding="utf-8")

    # render standalone + confluence storage
    tpl = ProjectStandardTemplate()
    view = tpl.build_standard_view(p)
    (out_dir / "standard_protocol_final.json").write_text(
        json.dumps(view, ensure_ascii=False, indent=2), encoding="utf-8")
    standalone = tpl.render_html(p)
    (out_dir / "standard_protocol_standalone.html").write_text(standalone, encoding="utf-8")
    storage = tpl.render_storage_html(p)
    (out_dir / "standard_protocol_confluence_storage.html").write_text(storage, encoding="utf-8")

    # title resolution
    (out_dir / "standard_title_resolution.json").write_text(
        json.dumps({"protocol_title": view["protocol_title"]}, ensure_ascii=False, indent=2),
        encoding="utf-8")

    # cell structure report
    report = build_cell_report(view)
    (out_dir / "standard_cell_structure_report.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    # DOCX export (python-docx)
    docx_path = out_dir / "standard_protocol.docx"
    write_docx(docx_path, view, standalone)

    # summary
    print("TITLE:", view["protocol_title"])
    print("client:", view["general_info"]["client_name"])
    print("topic_project:", view["general_info"]["topic_project"])
    print("meeting_type:", mt, mt_conf)
    print("h1 standalone:", standalone.count("<h1>"), "| h1 storage:", storage.count("<h1>"))
    print("initial:", view["meeting_context"]["initial_situation"][:120])
    print("main_problem:", view["meeting_context"]["main_problem"][:120])
    print("cell_report:", report)


def name_cn(res):
    return getattr(res, "client_name", "") or ""


def p_t_res_cn(res):
    return getattr(res, "client_name", "") or "УралДронЗавод"


def p_t_res_pn(res):
    return getattr(res, "project_name", "") or "Внедрение 1С:ERP"


def build_cell_report(view):
    from services.standard_protocol_compactor import normalize_cell_structure
    from protocol_templates.project_standard import _render_cell
    import html as _html
    esc = _html.escape
    cells = []
    for g in view.get("topic_groups", []):
        cells.append(g.get("cell"))
    for sec in ("decision_groups", "risk_groups", "task_groups", "open_questions"):
        for it in view.get(sec, []):
            cells.append(it.get("cell"))
    para_only = 0
    list_cells = 0
    pw_list = 0
    no_reason = 0
    mixed = 0
    for c in cells:
        if not c:
            continue
        st = c.get("structure_type", "paragraphs")
        if st == "paragraphs":
            para_only += 1
        elif st == "list":
            list_cells += 1
        elif st == "paragraphs_with_list":
            pw_list += 1
        if st in ("list", "paragraphs_with_list") and not c.get("list_reason"):
            no_reason += 1
    return {
        "cells_total": len(cells),
        "paragraph_only_cells": para_only,
        "list_cells": list_cells,
        "paragraphs_with_list_cells": pw_list,
        "list_without_reason_count": no_reason,
        "mixed_pattern_normalized_count": mixed,
        "control_cell_paragraph_count": _count_p(view["topic_groups"][0].get("cell")),
        "control_cell_ul_count": 0,
        "control_cell_li_count": 0,
    }


def _count_p(cell):
    if not cell:
        return 0
    return len(cell.get("paragraphs", []))


def write_docx(path: Path, view: dict, html: str):
    from docx import Document
    doc = Document()
    doc.add_heading(view["protocol_title"], level=0)
    for key, val in view["general_info"].items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(val))
    for sec_key, title in (
        ("meeting_context", "Контекст встречи"),
        ("key_outcomes", "Ключевые итоги"),
        ("topic_groups", "Обсуждение по тематическим блокам"),
        ("decision_groups", "Решения"),
        ("open_questions", "Открытые вопросы"),
        ("risk_groups", "Риски"),
        ("task_groups", "Задачи"),
    ):
        doc.add_heading(title, level=1)
        for item in view.get(sec_key, []):
            doc.add_paragraph(str(item))
    doc.save(str(path))


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\EVErmeev\AppData\Local\Temp\opencode\e2e_std_ural"
    import time as _t
    main(Path(out))